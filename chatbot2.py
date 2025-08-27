from flask import Flask, jsonify, request, render_template, send_file
from flask_cors import CORS
import random
import docx
import json
from datetime import datetime
from docx import Document
import PyPDF2
from werkzeug.utils import secure_filename
import os
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# === NEW IMPORTS ===
from search_engine import parse_query, hybrid_search

app = Flask(__name__)
CORS(app)

PAPERS_FILE = "saved_papers.json"
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER


# ===== Helper functions for user-created papers =====
def load_user_papers():
    if os.path.exists(PAPERS_FILE):
        with open(PAPERS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []


def save_user_papers(data):
    with open(PAPERS_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/create-paper")
def create_paper_page():
    return render_template("create_paper.html")


@app.route("/download-paper", methods=["POST"])
def download_paper():
    data = request.json
    format_type = data.get("format", "docx")

    if format_type == "docx":
        doc = Document()
        doc.add_heading(data.get("title", "Untitled Paper"), 0)
        if data.get("authors"):
            doc.add_paragraph(f"Authors: {data['authors']}")
            doc.add_paragraph("")

        for sec in data.get("sections", []):
            doc.add_heading(sec["title"], level=1)
            doc.add_paragraph(sec["content"])

        filepath = "generated_paper.docx"
        doc.save(filepath)

    elif format_type == "pdf":
        filepath = "generated_paper.pdf"
        doc = SimpleDocTemplate(filepath)
        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph(data.get("title", "Untitled Paper"), styles["Title"]))
        if data.get("authors"):
            story.append(Paragraph(f"Authors: {data['authors']}", styles["Normal"]))
        story.append(Spacer(1, 12))

        for sec in data.get("sections", []):
            story.append(Paragraph(sec["title"], styles["Heading2"]))
            story.append(Paragraph(sec["content"], styles["Normal"]))
            story.append(Spacer(1, 12))

        doc.build(story)

    return send_file(filepath, as_attachment=True)


@app.route("/save-paper", methods=["POST"])
def save_paper():
    paper = request.json
    if not paper:
        return jsonify({"message": "Invalid paper data"}), 400

    papers = load_user_papers()

    paper_entry = {
        "id": len(papers) + 1,
        "title": paper.get("title", "Untitled Paper"),
        "sections": paper.get("sections", []),
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }

    papers.append(paper_entry)
    save_user_papers(papers)

    return jsonify({"message": f"Paper '{paper_entry['title']}' saved successfully!"})


@app.route("/get-papers", methods=["GET"])
def get_papers():
    papers = load_user_papers()
    return jsonify(papers)


@app.route("/myspace")
def myspace_page():
    return render_template("myspace.html")


# ====== 🔑 MAIN CHAT ENDPOINT (uses hybrid search engine) ======
@app.route("/chat", methods=["POST"])
def chat():
    data = request.get_json()
    question = data.get("question", "").strip()

    if not question:
        return jsonify({"answer": "Please enter a valid question."}), 400

    filters = parse_query(question)
    results = hybrid_search(filters, question)

    if results:
        formatted = []
        for paper in results:
            link = paper.get("doi") or paper.get("alt_link") or "N/A"
            formatted.append(
                f"""
                <strong>Title:</strong> {paper.get('title','N/A')}<br>
                <strong>Authors:</strong> {paper.get('first_author','N/A')}<br>
                <strong>Journal:</strong> {paper.get('journal','N/A')}<br>
                <strong>Year:</strong> {paper.get('year','N/A')}<br>
                <strong>Score:</strong> {paper.get('_score',0)}<br>
                <strong>Link:</strong> <a href="{link}" target="_blank">{link}</a><br><br>
                """
            )
        return jsonify({"answer": "".join(formatted)})

    return jsonify({"answer": "No matching papers found."})


# ====== Plagiarism Check (simulated) ======
@app.route("/plagirism")
def plagiarism_page():
    return render_template("plagirism.html")


@app.route("/check-plagiarism", methods=["POST"])
def check_plagiarism():
    if "file" not in request.files:
        return jsonify({"message": "No file uploaded"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"message": "No selected file"}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(filepath)

    text = ""
    if filename.endswith(".txt"):
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
    elif filename.endswith(".docx"):
        doc = docx.Document(filepath)
        text = "\n".join([para.text for para in doc.paragraphs])
    elif filename.endswith(".pdf"):
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""

    plagiarism_percentage = random.randint(5, 75)
    return jsonify({"plagiarism": plagiarism_percentage})


if __name__ == "__main__":
    # Ensure templates exist (index.html, create_paper.html, plagirism.html)
    app.run(debug=True, port=5000)
